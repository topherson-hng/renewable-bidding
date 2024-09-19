import openpyxl
import logging
from openpyxl.styles import PatternFill
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
import datetime as dt
from docx import Document
import csv

def highlight_month_dates(file_path, sheet_name, target_sheet_name="HighlightedData"):
    # Load the workbook and the specified sheet
    wb = openpyxl.load_workbook(file_path)
    ws = wb['CurrentlyBidContracts']
    
    # checks if there is sheet already, if not creates a new sheet
    if target_sheet_name in wb.sheetnames:
            ws_target = wb[target_sheet_name]
    else:
        ws_target = wb.create_sheet(target_sheet_name)
    
    # adding headers to sheet
    ws_target.delete_rows(1, ws_target.max_row)
    ws_target['A1'] = "Log No"
    ws_target['B1'] = "Project Name"
    ws_target['C1'] = "Resource ID"
    ws_target['D1'] = "New Curtailment Date"

    # intialize target row
    target_row = 2
    word_row = 2

    # Define the fill style for highlighting
    highlight_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

    # Get the current date, month, and year
    today = dt.datetime.today()
    str_today = str(today)
    date_today = str_today.replace("-","")
    new_date_today = date_today.split(' ',1)[0]
    yyyymmdd = new_date_today.replace(" ","")
    current_month = today.month
    current_year = today.year
    
    # create word document
    test_document = Document()
    first_sentence = test_document.add_paragraph("Curtailment Anniversaries:")
    first_sentence.bold = True

    # Iterate through the cells in column L
    for row in ws.iter_rows(min_col=12, max_col=12, min_row=1, max_row=ws.max_row):
        for cell in row:
            if isinstance(cell.value, dt.datetime):
                date_value = cell.value
                
                if date_value.month != current_month:
                    cell.fill = PatternFill()  # Clear the highlight
                elif date_value.month == current_month and date_value.year == current_year:
                    cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
                    cell.value = date_value.replace(year=date_value.year + 1)
                   
            # pull highlighted cell value row "L" into separate sheet    
            if cell.fill == PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid"):
              # Pull the values from columns A, B, and C
                ws_target.cell(row=target_row, column=1).value = ws.cell(row=cell.row, column=3).value
                ws_target.cell(row=target_row, column=2).value = ws.cell(row=cell.row, column=2).value
                ws_target.cell(row=target_row, column=3).value = ws.cell(row=cell.row, column=5).value
                ws_target.cell(row=target_row, column=4).value = ws.cell(row=cell.row, column=12).value
                target_row += 1

    for row in ws_target.iter_rows(min_col=4, max_col=4, min_row=1, max_row=ws_target.max_row):
        for cell in row:
            cell.number_format = 'mm/dd/yyyy;@'


    max_row_for_d = max((d.row for d in ws_target['D'] if d.value is not None))
    style = TableStyleInfo(name="TableStyleMedium9", showRowStripes=True)
    table = Table(displayName="table1", ref="A1:" + get_column_letter(ws_target.max_column) + str(max_row_for_d))
    table.tableStyleInfo = style
    ws_target.add_table(table)

    # import specified cells into word file 
    for row in ws_target.iter_rows(min_row=2, max_row=ws_target.max_row, min_col=1, max_col=ws_target.max_column):
            log_cell = ws_target.cell(row=word_row, column=1).value
            name_cell = ws_target.cell(row=word_row, column=2).value
            resource_cell = ws_target.cell(row=word_row, column=3).value
            date_cell = ws_target.cell(row=word_row, column=4).value
            # dt_cell = date_cell.ToString("MM/dd/yyyy", CultureInfo.InvariantCulture)
            
            if row[0].value:
                test_document.add_paragraph(f"{name_cell} - Resource ID: {resource_cell}")
                test_document.add_paragraph(f"The new curtailment anniversary is on {date_cell}")
                word_row += 1

    # Save the workbook 
    test_document.save(f'C:/Users/cyh3/OneDrive - PGE/Desktop/test/Curtailment Info {yyyymmdd}.docx')
    wb.save(f'C:/Users/cyh3/OneDrive - PGE/Desktop/test/Renewable Bidding - Upcoming Contracts {yyyymmdd}.xlsx')

# Usage 
file_path = r'C:/Users/cyh3/OneDrive - PGE/Desktop/test/Renewable Bidding - Upcoming Contracts 20240822.xlsx'
sheet_name = 'CurrentlyBidContracts'
highlight_month_dates(file_path, sheet_name)
